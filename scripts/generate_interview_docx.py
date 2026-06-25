from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = r"D:\develop\codex-project\project-001-love-travel-map\interview-questions-working.docx"


def set_run_font(run, name="Microsoft YaHei", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_run_font(run, size=22, bold=True, color="1F4D78")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle)
    set_run_font(run, size=10, color="666666")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True, color="2E74B5" if level == 1 else "1F4D78")


def add_label_paragraph(doc, label, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    label_run = p.add_run(label)
    set_run_font(label_run, bold=True, color="1F4D78")
    body_run = p.add_run(body)
    set_run_font(body_run)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run_font(run)


def add_question(doc, title, concept, project_use, why, interview_points):
    add_heading(doc, title, level=2)
    add_label_paragraph(doc, "概念：", concept)
    add_label_paragraph(doc, "项目中怎么用：", project_use)
    add_label_paragraph(doc, "为什么这样用：", why)
    add_label_paragraph(doc, "面试回答要点：", "")
    add_bullets(doc, interview_points)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    add_title(
        doc,
        "情侣旅行地图项目面试题库",
        "记录项目中可以在面试里讲清楚的技术点：概念、项目落地方式、为什么这样设计、回答重点。"
    )

    add_heading(doc, "一、Redis 在项目中的使用", level=1)

    add_question(
        doc,
        "1. Redis 在项目里主要解决什么问题？",
        "Redis 是基于内存的高性能 key-value 数据库，常用于缓存、分布式锁、限流、验证码、Session、排行榜、消息队列辅助等短期高频场景。",
        "本项目没有把 Redis 当成 MySQL 的替代品，而是用于保存短期状态：登录 Session 白名单、邀请码 1 分钟有效期、AI 规划防重复点击锁、AI 规划短窗口限流。",
        "旅行日记、图片、计划、空间成员这些长期业务数据必须落 MySQL，因为它们需要持久化、事务和可追溯。Redis 更适合处理短期状态和高频访问，能降低数据库压力，也能增强安全控制。",
        [
            "我把长期数据和短期状态分开：MySQL 保存核心业务数据，Redis 保存会过期的状态。",
            "Redis 在项目里承担 Session、邀请码 TTL、防重复点击和限流，而不是盲目缓存所有数据。",
            "这样既能体现 Redis 能力，也避免缓存和数据库一致性变复杂。"
        ]
    )

    add_question(
        doc,
        "2. 为什么登录 Session 要放 Redis？",
        "Session 是服务端保存用户登录状态的一种机制。传统方式可以保存在服务器内存，但多实例部署时内存不共享，用户请求打到不同机器会出现登录状态不一致。",
        "本项目使用 HttpOnly Cookie 保存签名后的会话标识，同时在 Redis 中保存 `love-travel:session:{hash}`。用户请求接口时，后端不仅校验 Cookie 签名，还必须在 Redis 查到对应 Session，才认为登录有效。",
        "这样做可以支持退出登录立即失效、后续多后端实例共享登录状态，也避免只靠前端或纯 Cookie 判断登录。HttpOnly Cookie 防止前端 JS 直接读取，Redis 白名单让后端掌握是否允许继续登录。",
        [
            "Cookie 负责携带登录凭证，Redis 负责服务端确认这个凭证是否仍然有效。",
            "退出登录时删除 Redis Session，即使 Cookie 还在，也不能继续访问接口。",
            "这种方式比单机内存 Session 更适合企业里的多实例部署。"
        ]
    )

    add_question(
        doc,
        "3. Redis Session 如何避免缓存雪崩？",
        "缓存雪崩是大量缓存 key 在同一时间集中失效，导致请求瞬间打到后端数据库或认证逻辑，引发压力峰值。",
        "本项目 Session 的 Redis TTL 是 7 天，并额外增加随机抖动。也就是不同用户的 Session 不会精确在同一秒过期。",
        "登录 Session 数量可能随着用户增长变多，如果大量 Session 同时过期，会导致用户集中重新登录或接口集中失败。TTL 随机抖动能把失效时间打散。",
        [
            "我没有让所有 Session 统一固定秒数过期，而是在 TTL 上加随机值。",
            "随机过期可以把同一时间的失效压力摊开。",
            "这就是项目里针对缓存雪崩做的实际处理。"
        ]
    )

    add_question(
        doc,
        "4. 登录失败次数限制为什么适合用 Redis？",
        "登录失败次数限制是防暴力破解的一种常见安全策略。它需要记录某个账号在一段时间内连续失败了多少次，并在达到阈值后临时锁定。",
        "本项目使用 Redis 保存 `love-travel:login:fail:{account}` 和 `love-travel:login:lock:{account}`。10 分钟内密码输错 5 次，就写入锁定 key，锁定 10 分钟。登录成功后清理失败计数和锁定状态。",
        "这个状态是短期的、自动过期的，非常适合 Redis。如果放 MySQL，需要频繁写库，还要额外清理过期记录；Redis 的自增计数和 TTL 更轻量。",
        [
            "我用 Redis 计数器记录登录失败次数，用 TTL 自动清理。",
            "达到阈值后写入锁定 key，避免暴力破解密码。",
            "提示语不直接暴露账号是否存在，降低账号枚举风险。"
        ]
    )

    add_question(
        doc,
        "5. 邀请码为什么适合放 Redis？",
        "邀请码、验证码这类数据有明显的短生命周期，例如 1 分钟或 5 分钟有效。Redis 天然支持 TTL，适合保存这种自动过期的数据。",
        "本项目的情侣空间邀请码生成后，会写入 Redis：`love-travel:invite-code:{code}`，TTL 约 1 分钟。用户输入邀请码加入空间时，后端先查 Redis。如果 Redis 中不存在，就认为邀请码不存在或已过期。",
        "如果只靠 MySQL 定时清理过期邀请码，需要额外任务；使用 Redis TTL 能自然过期。MySQL 仍然保存邀请码记录，用于审计和排查问题。",
        [
            "邀请码是典型短期状态，Redis TTL 比 MySQL 定时清理更合适。",
            "Redis 负责有效期和快速判断，MySQL 负责历史记录。",
            "这种设计兼顾性能、自动过期和可追溯。"
        ]
    )

    add_question(
        doc,
        "6. AI 规划为什么要做防重复点击？",
        "防重复点击本质上是防止同一个用户在短时间内提交多次相同或冲突请求。对于调用大模型的接口，这不仅影响数据正确性，还会造成真实成本浪费。",
        "本项目在用户生成某一天 AI 旅行计划时，使用 Redis `SET NX EX` 创建锁：`love-travel:ai:plan-day:generating:{userId}:{dayId}`。同一用户同一天如果已有生成任务，后端直接拒绝新的生成请求。",
        "AI 规划会调用 Python 服务和通义千问模型，成本比普通 CRUD 高。如果用户连点按钮，可能重复扣额度、重复调用模型、产生多个草稿。Redis 锁可以在 Java 后端入口处提前挡住。",
        [
            "`SET NX` 保证只有第一个请求能拿到锁。",
            "`EX` 设置过期时间，避免服务异常导致锁永久不释放。",
            "生成完成或失败后会主动删除锁，用户可以再次生成。"
        ]
    )

    add_question(
        doc,
        "7. AI 规划限流怎么做？",
        "限流是限制某个用户或某个接口在单位时间内的请求次数，防止接口被误触、刷接口或恶意调用。",
        "本项目为 AI 规划增加了 Redis 短窗口限流 key：`love-travel:ai:plan-day:rate:{userId}`。每次请求递增计数，第一次请求设置 1 分钟过期，超过阈值就返回“请求太频繁”。",
        "AI 规划接口比普通接口更贵，会消耗模型额度和服务器资源。短窗口限流可以保护 Java 后端、Python AI 服务和大模型 API。",
        [
            "我用 Redis 的自增计数实现用户级短窗口限流。",
            "第一次访问设置过期时间，窗口结束后计数自动清空。",
            "限流和 30 天额度不同：限流防短时间冲击，额度控制长期成本。"
        ]
    )

    add_question(
        doc,
        "8. 项目里如何理解缓存穿透？怎么避免？",
        "缓存穿透是指大量请求查询缓存和数据库中都不存在的数据，导致请求绕过缓存直接打到数据库。",
        "本项目没有把旅行记录、空间、计划这些长期对象做成简单缓存，所以没有引入大量对象缓存穿透风险。Session 会先校验 Cookie 签名，非法 Cookie 不会继续查 Redis 或数据库；邀请码 Redis miss 后只走受控的过期处理逻辑。",
        "对初版项目来说，不盲目缓存业务对象反而更安全。需要缓存时，要考虑空值缓存、布隆过滤器、参数校验等手段。",
        [
            "我先通过参数校验和签名校验过滤非法请求。",
            "没有为了炫技把所有 MySQL 查询都加缓存，避免引入穿透和一致性问题。",
            "如果后续缓存城市记录等热点数据，可以加短期空值缓存或布隆过滤器。"
        ]
    )

    add_question(
        doc,
        "9. 项目里如何理解缓存击穿？怎么避免？",
        "缓存击穿是某个热点 key 突然失效，大量请求同时访问这个 key，导致后端瞬间承压。也可以理解为并发请求同时穿过缓存去重建同一份数据。",
        "本项目 AI 规划用 Redis 生成锁解决类似问题：同一用户同一天的计划生成，只允许一个请求进入真正的 AI 调用流程，其他重复请求直接被拒绝。",
        "虽然这里不是传统热点商品缓存，但“多个请求同时触发同一个昂贵操作”的问题本质类似。用互斥锁可以避免昂贵资源被并发击穿。",
        [
            "我用 Redis 锁防止同一份 AI 计划被并发重复生成。",
            "锁有 TTL，防止异常情况下死锁。",
            "这属于把缓存击穿思想应用到高成本 AI 接口保护上。"
        ]
    )

    add_question(
        doc,
        "10. Redis 和 MySQL 在项目中的分工是什么？",
        "MySQL 适合保存强一致、长期存在、需要事务和关系查询的数据；Redis 适合保存短期、高频、自动过期、可快速判断的状态。",
        "本项目 MySQL 保存用户、空间、成员、旅行日记、照片记录、计划、AI 运行记录。Redis 保存 Session、邀请码有效期、AI 防重复锁和限流计数。",
        "这样可以减少 Redis 和 MySQL 的一致性压力。Redis key 即使过期，也不会丢失核心业务数据；MySQL 保留完整业务事实。",
        [
            "核心数据必须落 MySQL，Redis 不做唯一真相来源。",
            "Redis 只保存可以过期、可以重建、或用于短期控制的数据。",
            "这种分工更符合企业项目的常见架构。"
        ]
    )

    add_question(
        doc,
        "11. 如果面试官问 Redis 这块有什么可扩展点，怎么回答？",
        "可扩展点是指当前项目已经有基础能力，后续可以顺着同一技术方向继续增强，而不是孤立 demo。",
        "本项目后续可以继续用 Redis 做登录失败次数限制、接口级 IP 限流、热点地图数据缓存、AI 任务队列状态、WebSocket 在线状态、空间成员在线提醒等。",
        "这些扩展都围绕项目真实业务，不是为了堆技术名词。比如 AI 任务变慢后，可以把任务状态写 Redis，前端轮询或 SSE 获取进度。",
        [
            "我会先讲当前已落地的 Session、邀请码、AI 锁和限流。",
            "再讲后续可以扩展到登录失败限制、热点缓存、在线状态、AI 任务状态。",
            "重点强调技术是围绕业务问题引入的，不是为了简历堆栈。"
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("后续维护规则：以后每新增一个适合面试表达的技术点，都按“题目、概念、项目中怎么用、为什么这样用、面试回答要点”的结构追加到本文档。")
    set_run_font(run, size=10, color="666666")

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
